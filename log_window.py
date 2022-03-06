import base64
import datetime
import json
import shutil
import time
import tkinter
import tkinter as tk
from tkinter import *
from docx import Document
from tkinter import ttk
import re
import tkinter.messagebox
from tkinter import filedialog
from tkcalendar import Calendar, DateEntry
import win32api
import os
from threading import Thread
import xlwt
import docx
from docx import Document  # 用来建立一个word对象
from docx.shared import Pt  # 用来设置字体的大小
from docx.oxml.ns import qn  # 设置字体
from docx.shared import RGBColor  # 设置字体的颜色

#翻译
# if "804" not in "%x" % win32api.GetSystemDefaultLangID():
if False:
    language = "zh"
    start_date_des = "  起始日期："
    start_time_des = "起始时间："
    end_date_des = "截止日期："
    end_time_des = "截止时间："
    specific_user = "  指定用户："
    process_type = "流程类型："
    search_type = "搜索类型："
    chosen_type = "选择类型："
    process_des1 = "流程化模块"
    process_des2 = "非流程化模块"
    menu_des1 = "用户菜单"
    menu_des2 = "角色菜单"
    menu_des3 = "组织菜单"
    menu_des4 = "公司菜单"
    menu_des5 = "部门菜单"
    select1 = {"流程化模块": ("新建流程", "更改流程", "删除流程", "运行流程", "放行流程", "定时流程", "指挥流程", "所有"),
               "非流程化模块": ("用户菜单", "角色菜单", "组织菜单", "公司菜单", "部门菜单", "所有"),
               "":()}
    select2 = {"用户菜单": ("新建用户", "更改用户", "禁用/启用用户", "用户权限设置", "设置用户授权期限", "重置用户密码", "用户分配角色", "删除用户"),
               "角色菜单": ("新建角色", "更改角色", "角色权限设置", "删除角色", "设置角色授权期限", "角色分配用户"),
               "组织菜单": ("新建组织", "描述更改", "禁用/启用组织", "删除组织"),
               "公司菜单": ("新建公司", "描述更改", "禁用/启用公司", "删除公司"),
               "部门菜单": ("新建部门", "描述更改", "禁用/启用部门", "删除部门"),
               "新建流程": ("",),
               "更改流程": ("",),
               "运行流程": ("",),
               "放行流程": ("",),
               "定时流程": ("",),
               "指挥流程": ("",),
               "清理日志": ("",),
               "删除流程": ("",)}
    btn1_des = "执行"
    btn2_des = "退出"
    btn2_message = "是否确认退出？"
    time_error1 = "开始时间请小于结束时间！"
    time_error2 = "请输入正确的时间格式！"
    path_error = "请确保工具在RPA路径下！"
    operator_des = "操作者："
    operate_time_des = "运行时间："
    run_des1 = "文件路径："
    run_des2 = "运行流程节点："
    run_des3 = "运行结果："
    run_des4 = "失败步骤："
    run_des5 = "失败节点："
    run_des6 = "错误信息："
    save_des1 = "保存文件路径："
    save_des2 = "流程节点："
    save_des3 = "保存结果："
    update_des1 = "保存文件路径："
    update_des2 = "保存前流程节点："
    update_des3 = "保存后流程节点："
    update_des4 = "保存结果："
    delete_des1 = "删除文件路径："
    delete_des2 = "流程节点："
    delete_des3 = "删除结果："
    release_des1 = "原始路径："
    release_des2 = "放行路径："
    booked_des1 = "定时时间："
    booked_des2 = "流程路径："
    booked_des3 = "流程明细："
    booked_des4 = "被执行的机器人："
    user_des = "用户信息："
    change_user_des1 = "更改前用户信息："
    change_user_des2 = "更改后用户信息："
    username_des = "用户名："
    user_enable_disable_des2 = "操作动作（禁用/启用）："
    user_permission_setting_des2 = "更改前权限清单："
    user_permission_setting_des3 = "更改后权限清单："
    user_authorization_period_des2 = "期限值："
    user_assigned_roles_des2 = "分配的角色："
    role_update_des1 = "更改前角色信息："
    role_update_des2 = "更改后角色信息："
    userlist_des = "用户清单："
    detail_des = "明细："
    update_des1 = "更改前描述："
    update_des2 = "更改后描述："
    organization_des = "组织名："
    company_des = "公司名："
    department_des = "部门名："
else:
    language = "en"
    start_date_des = "  Start Date:"
    start_time_des = "Start Time:"
    end_date_des = "End Date:"
    end_time_des = "End Time:"
    specific_user = "  User:"
    process_type = "Process:"
    search_type = "Menu:"
    chosen_type = "Function:"
    process_des1 = "Process"
    process_des2 = "Non-process"
    menu_des1 = "User"
    menu_des2 = "Role"
    menu_des3 = "Organization"
    menu_des4 = "Company"
    menu_des5 = "Department"
    select1 = {"Process": ("New", "Update", "Delete", "Run", "Release", "Booking", "Commander", "All"),
                     "Non-process": ("User", "Role", "Organization", "Company", "Department", "All"),
                     "":""}
    select2 = {"User": ("New", "Update", "Active/Inactive", "Delete", "Permission", "Indate", "Reset Password", "Assign Role", "Delete"),
                   "Role": ("New", "Update", "Permission", "Delete", "Permission", "Assign User"),
                   "Organization": ("New", "Update", "Active/Inactive", "Delete"),
                   "Company": ("New", "Update", "Active/Inactive", "Delete"),
                   "Department": ("New", "Update", "Active/Inactive", "Delete"),
                   "New": ("",),
                   "Delete": ("",),
                   "Update": ("",),
                   "Run": ("",),
                   "Release": ("",),
                   "Booking": ("",),
                   "Commander": ("",)}
    btn1_des = "Execute"
    btn2_des = "Quit"
    btn2_message = "Are you sure to exit?"
    time_error1 = "Start time muse be after the end time!"
    time_error2 = "Please input the correct time format!"
    path_error = "Please make sure the tool is in the RPA path!"
    operator_des = "Operator:"
    operate_time_des = "Operate Time:"
    run_des1 = "Process Path:"
    run_des2 = "Process Nodes:"
    run_des3 = "Result:"
    run_des4 = "Error Step:"
    run_des5 = "Error Node:"
    run_des6 = "Error Message:"
    save_des1 = "Saved Path:"
    save_des2 = "Process Nodes:"
    save_des3 = "Saved Result:"
    update_des1 = "Saved Path:"
    update_des2 = "Process Nodes Before Saved:"
    update_des3 = "Process Nodes After Saved:"
    update_des4 = "Saved Result:"
    delete_des1 = "Deleted Path:"
    delete_des2 = "Process Nodes:"
    delete_des3 = "Delete Result:"
    release_des1 = "Original Path:"
    release_des2 = "Current Path:"
    booked_des1 = "Booking Time:"
    booked_des2 = "Booking Process:"
    booked_des3 = "Process Nodes:"
    booked_des4 = "Executed Robot:"
    user_des = "User Information:"
    change_user_des1 = "User Information Before Change:"
    change_user_des2 = "User Information After Change:"
    username_des = "Username:"
    user_enable_disable_des2 = "Operation(Active/Inactive):"
    user_permission_setting_des2 = "Permissions Before Update:"
    user_permission_setting_des3 = "Permissions After Update:"
    user_assigned_roles_des2 = "Assigned Roles:"
    role_update_des1 = "Role Information Before Update:"
    role_update_des2 = "Role Information After Update:"
    userlist_des = "User List:"
    detail_des = "Detail:"
    update_des1 = "Description Before Update:"
    update_des2 = "Description After Update:"
    organization_des = "Organization:"
    company_des = "Company:"
    department_des = "Department:"

log_foler = ""

def func_process(process):
    """
    :param process: 字典嵌套字典的json文件
    :return: 流程节点数据处理
    """
    process_dict = json.loads(process)
    process_data = ""
    for k, i in process_dict.items():
        process_func = [j for j in i.keys()][0]
        process_parameter = [j for j in i.values()][0]
        arg_data = ""
        if type(process_parameter) != dict:
            process_parameter = json.loads(process_parameter)
        for k_name,arg in process_parameter.items():
            if k_name == 'divId' or k_name == 'groupId'or k_name == "process_background_color":
                continue
            if k_name == "if_condition":
                arg_data = arg_data + f"{k_name}({arg})" + ','
            elif k_name == "self_variant":
                arg_data = arg_data + f"{arg}" + ','
            else:
                arg_data = arg_data + f"{k_name}={arg}" + ','
        # placeholder = ''
        # process_data = process_data + placeholder + f"# node_no:{k} "+ process_func +'\n'+ arg_data + '\n\n'
        process_data = process_data + f"# node_no:{k} "+ process_func +'\n'+ arg_data + '\n'
        # if k_name == "If" or k_name == "For" or k_name == "Else":
        #     placeholder += "    "
        # elif k_name == "End_If" or k_name == "Exit_For":
        #     placeholder -= "    "
    # process_data = process_data + "End"
    return process_data

def dateentry_view(window):
    """
    时间选择器
    :param window: window = tk.Tk()
    :return: 年月日
    """
    def print_sel(e):
        print(cal.get_date())
    top = tk.Toplevel(window)

    ttk.Label(top, text='Choose date').pack(padx=10, pady=10)
    cal = DateEntry(top, width=12, background='CornflowerBlue',
                    foreground='white', borderwidth=2)
    cal.pack(padx=10, pady=10)
    cal.bind("<<DateEntrySelected>>", print_sel)
    return cal.get_date()

#解析log文件
def read_log_file(user, path, action_code, start_time, end_time):
    with open(path, 'r', encoding='utf-8') as f:
        data_list = f.readlines()
    result_list = []
    for num, line_data in enumerate(data_list):
        try:
            time_str = line_data[1:20]
            time_str_2_date = datetime.datetime.strptime(time_str, "%Y-%m-%d %H:%M:%S")
            if time_str_2_date >= start_time and time_str_2_date <= end_time and ">{" in line_data:
                if action_code == "100":
                    try:
                        action_code_no = str(re.search("- (.*)>{", line_data, re.M | re.I).group()).replace("- ","").replace(">{","")
                        if action_code_no not in ["415","425","345","200","360","400","500","600","601","602","603","604","605","606","700","701","702","703","705","704","800","801", "802","803","900","901","902","903","1000","1001","1002", "1003"]:
                            continue
                    except:
                        pass
                elif action_code == "101":
                    try:
                        action_code_no = str(re.search("- (.*)>{", line_data, re.M | re.I).group()).replace("- ","").replace(">{", "")
                        if action_code_no not in ["415", "425", "345", "200", "360", "400", "500", "600"]:
                            continue
                    except:
                        pass
                elif action_code == "102":
                    try:
                        action_code_no = str(re.search("- (.*)>{", line_data, re.M | re.I).group()).replace("- ","").replace(">{", "")
                        if action_code_no not in ["600","601","602","603","604","605","606","700","701","702","703","705","704","800","801", "802","803","900","901","902","903","1000","1001","1002", "1003"]:
                            continue
                    except:
                        pass
                else:
                    action_code_no = action_code
                if action_code_no in line_data:
                    result = str(re.search(r"{}>(.*)".format(action_code_no), line_data, re.M | re.I).group())
                    if user:
                        result_json = json.loads(result.replace("{}>".format(action_code_no), ""))
                        if "username" in result_json:
                            if result_json["username"].lower() == user.lower():
                                result_list.append(result.replace("{}>".format(action_code_no), ""))
                    else:
                        result_list.append(result.replace("{}>".format(action_code_no), ""))
            elif time_str_2_date > end_time:
                break
        except Exception:
            pass
    return {path:result_list}

class MyThread(Thread):
    def __init__(self,user,number,length,file_list,action_code,start_time,end_time):
        Thread.__init__(self)
        self.user = user
        self.number = number
        self.length = length
        self.file_list = file_list
        self.action_code = action_code
        self.start_time = start_time
        self.end_time = end_time

    def read_file(self):
        result = {}
        for i in range(self.number, self.length, 20):
            try:
                file_path = self.file_list[i]
                file_result = read_log_file(self.user, file_path, self.action_code, self.start_time, self.end_time)
                result = dict(result, **file_result)
            except:
                break
        return result

    def run(self):
        self.result = self.read_file()

    def get_result(self):
        return self.result

def log_window():
    global log_foler
    """日志分析器"""
    def to_exit():
        """退出确定"""
        exit = tkinter.messagebox.askokcancel(title=btn2_des, message=btn2_message)
        if exit:
            window.destroy()

    def text_contract():
        """解析合同内容"""
        document = Document('TecleadContract.docx')
        contract_data = ''
        for i in document.paragraphs:
            contract_data = f'{contract_data}\n{i.text}'
        return contract_data

    def log_message(method):
        """在log中读取信息"""
        # initial_time:初始时间。last_time：截止时间。username：用户名
        try:
            initial_time = str(e1.get()) + " " + str(e3.get())
            last_time = str(e2.get()) + " " + str(e4.get())
            try:
                initial_time_2_date = datetime.datetime.strptime(initial_time, "%Y-%m-%d %H:%M:%S")
                last_time_2_date = datetime.datetime.strptime(last_time, "%Y-%m-%d %H:%M:%S")
                if initial_time_2_date > last_time_2_date:
                    tkinter.messagebox.showwarning(title="Error", message=time_error1)
                    return "Error"
            except Exception:
                tkinter.messagebox.showwarning(title="Error", message=time_error2)
                return "Error"
            username = e5.get()
            #获取所有符合日期条件的log文件
            file_list = []
            initial_date = datetime.datetime.strptime(str(e1.get()), "%Y-%m-%d")
            last_date = datetime.datetime.strptime(str(e2.get()), "%Y-%m-%d")
            next_day = initial_date + datetime.timedelta(days=-1)
            while True:
                next_day = next_day + datetime.timedelta(days=1)
                if next_day <= last_date:
                    log_file_path = log_foler + "\\info-" + next_day.strftime("%Y-%m-%d") + ".log"
                    if os.path.exists(log_file_path):
                        file_list.append(log_file_path)
                else:
                    break
            #动作参数列表
            action_json = {"all": "100",
                           "process_all": "101",
                           "non_process_all": "102",
                           "save": "415",
                           "update": "425",
                           "delete": "345",
                           "run": "200",
                           "release": "360",
                           "timing": "400",
                           "commander": "500",
                           "user_new": "600",
                           "user_update": "601",
                           "user_enable_disable": "602",
                           "user_permission_setting": "603",
                           "user_authorization_period": "604",
                           "user_reset_password": "605",
                           "user_assigned_roles": "606",
                           "user_delete": "607",
                           "role_new": "700",
                           "role_update": "701",
                           "role_permission_setting": "702",
                           "role_delete": "703",
                           "role_authorization_period": "705",
                           "role_assigned_roles": "704",
                           "organization_new": "800",
                           "organization_update": "801",
                           "organization_enable_disable": "802",
                           "organization_delete": "803",
                           "company_new": "900",
                           "company_update": "901",
                           "company_enable_disable": "902",
                           "company_delete": "903",
                           "department_new": "1000",
                           "department_update": "1001",
                           "department_enable_disable": "1002",
                           "department_delete": "1003",
                           }
            action_code = action_json[method]
            #todo
            #此处插入分线程函数（最多20个线程）
            thread_list = []
            result_json = {}
            length = len(file_list)
            if length > 20:
                fileQty = 20
            else:
                fileQty = length
            for i in range(fileQty):
                t = MyThread(username, i, length, file_list, action_code, initial_time_2_date, last_time_2_date)
                thread_list.append(t)
            for t in thread_list:
                t.start()
            for t in thread_list:
                t.join()
                thread_result = t.get_result()
                result_json = dict(result_json, **thread_result)
            #调整结果数组顺序
            result_list = []
            for file_path in file_list:
                file_result = result_json[file_path]
                if len(file_result) > 0:
                    result_list+=file_result
            #将结果整理成字符串
            log_display_content = ""
            for result in result_list:
                try:
                    result_json = json.loads(result)
                    for key,value in result_json.items():
                        if isinstance(value, list):
                            log_display_content = log_display_content + str(key).capitalize() + ":\n"
                            for val in value:
                                if isinstance(val, dict):
                                    log_display_content+=val["code"]
                                else:
                                    log_display_content += val
                        else:
                            log_display_content = log_display_content + str(key).capitalize() + ":\n" + str(value) + "\n"
                    log_display_content = log_display_content[0:len(log_display_content)-1]
                    log_display_content = log_display_content + '\n' + '=' * 200 + '\n' + '=' * 200 + "\n"
                except Exception:
                    pass
            if not log_display_content:
                log_display_content = "No Data!"
            return log_display_content
        except Exception:
            pass

    def boxlist(*args):  # 处理事件，*args表示可变参数
        """ boxlist日志展示"""
        mylist.delete(0, END)
        log_display_content = log_message(*args)
        if log_display_content:
            data_list = log_display_content.split('\n')
        color = "no"
        if log_display_content == 'Error':
            return False
        elif log_display_content == 'No Data':
            log_display_content = "No data!" + '\n' + '=' * 1000 + '\n' + '=' * 1000 + "\n"
            data_list = log_display_content.split('\n')
        if data_list:
            for num,data in enumerate(data_list):
                mylist.insert(END, data)
                if "#" in data:
                    color = "#808080"
                    mylist.itemconfig(END, fg="#808080")
                elif "====" in data:
                    mylist.itemconfig(END, fg="black")
                    color = "no"
                elif color == "#808080":
                    color = "no"
                elif color == "no":
                    if data == "Action:":
                        mylist.itemconfig(END, fg="black")
                    else:
                        mylist.itemconfig(END, fg="black")
                    color = "yes"
                elif color == "yes":
                    mylist.itemconfig(END, fg="#D2691E")
                    color = "no"
                # if (num % 2) == 0:
                #     mylist.itemconfig(END, fg="red")
            # 列表框位于窗口左端，当窗口改变

    def getting_click():
        """获取信息点击事件"""
        print("起始时间：%s" % e1.get())
        print("截至时间：%s" % e2.get())
        print("用户名：%s" % e3.get())
        print(comboxlist3.get())
        # 自動刪除
        # e1.delete(0, END)
        # e2.delete(0, END)
        # e3.delete(0, END)
        if comboxlist2.get() == select1[process_des1][0]:
            boxlist("save")
        elif comboxlist2.get() == select1[process_des1][1]:
            boxlist("update")
        elif comboxlist2.get() == select1[process_des1][2]:
            boxlist("delete")
        elif comboxlist2.get() == select1[process_des1][3]:
            boxlist("run")
        elif comboxlist2.get() == select1[process_des1][4]:
            boxlist("release")
        elif comboxlist2.get() == select1[process_des1][5]:
            boxlist("timing")
        elif comboxlist2.get() == select1[process_des1][6]:
            boxlist("commander")
        elif comboxlist2.get() == "All" and comboxlist1.get() == process_des1:
            boxlist("process_all")
        elif comboxlist2.get() == "All" and comboxlist1.get() == process_des2:
            boxlist("non_process_all")
        elif comboxlist2.get() == "":
            boxlist("all")
        # "用户菜单": ("新建用户", "更改用户", "禁用/启用用户", "用户权限设置", "设置用户授权期限", "重置用户密码", "用户分配角色"
        elif comboxlist2.get() == menu_des1 and comboxlist3.get() == select2[menu_des1][0]:
            boxlist("user_new")
        elif comboxlist2.get() == menu_des1 and comboxlist3.get() == select2[menu_des1][1]:
            boxlist("user_update")
        elif comboxlist2.get() == menu_des1 and comboxlist3.get() == select2[menu_des1][2]:
            boxlist("user_enable_disable")
        elif comboxlist2.get() == menu_des1 and comboxlist3.get() == select2[menu_des1][4]:
            boxlist("user_permission_setting")
        elif comboxlist2.get() == menu_des1 and comboxlist3.get() == select2[menu_des1][5]:
            boxlist("user_authorization_period")
        elif comboxlist2.get() == menu_des1 and comboxlist3.get() == select2[menu_des1][6]:
            boxlist("user_reset_password")
        elif comboxlist2.get() == menu_des1 and comboxlist3.get() == select2[menu_des1][7]:
            boxlist("user_assigned_roles")
        elif comboxlist2.get() == menu_des1 and comboxlist3.get() == select2[menu_des1][8]:
            boxlist("user_delete")
            # "角色菜单": ("新建角色", "更改角色", "角色权限设置", "删除角色", "设置角色授权期限", "角色分配用户"
        elif comboxlist2.get() == menu_des2 and comboxlist3.get() == select2[menu_des2][0]:
            boxlist("role_new")
        elif comboxlist2.get() == menu_des2 and comboxlist3.get() == select2[menu_des2][1]:
            boxlist("role_update")
        elif comboxlist2.get() == menu_des2 and comboxlist3.get() == select2[menu_des2][2]:
            boxlist("role_permission_setting")
        elif comboxlist2.get() == menu_des2 and comboxlist3.get() == select2[menu_des2][3]:
            boxlist("role_delete")
        elif comboxlist2.get() == menu_des2 and comboxlist3.get() == select2[menu_des2][4]:
            boxlist("role_authorization_period")
        elif comboxlist2.get() == menu_des2 and comboxlist3.get() == select2[menu_des2][5]:
            boxlist("role_assigned_roles")
            # "组织菜单": ("新建组织", "描述更改", "禁用/启用组织", "删除组织"
        elif comboxlist2.get() == menu_des3 and comboxlist3.get() == select2[menu_des3][0]:
            boxlist("organization_new")
        elif comboxlist2.get() == menu_des3 and comboxlist3.get() == select2[menu_des3][1]:
            boxlist("organization_update")
        elif comboxlist2.get() == menu_des3 and comboxlist3.get() == select2[menu_des3][2]:
            boxlist("organization_enable_disable")
        elif comboxlist2.get() == menu_des3 and comboxlist3.get() == select2[menu_des3][3]:
            boxlist("organization_delete")
            # "公司菜单": ("新建公司", "描述更改", "禁用/启用公司", "删除公司"
        elif comboxlist2.get() == menu_des4 and comboxlist3.get() == select2[menu_des4][0]:
            boxlist("company_new")
        elif comboxlist2.get() == menu_des4 and comboxlist3.get() == select2[menu_des4][1]:
            boxlist("company_update")
        elif comboxlist2.get() == menu_des4 and comboxlist3.get() == select2[menu_des4][2]:
            boxlist("company_enable_disable")
        elif comboxlist2.get() == menu_des4 and comboxlist3.get() == select2[menu_des4][3]:
            boxlist("company_delete")
            # "部门菜单": ("新建部门", "描述更改", "禁用/启用部门", "删除部门"
        elif comboxlist2.get() == menu_des5 and comboxlist3.get() == select2[menu_des5][0]:
            boxlist("department_new")
        elif comboxlist2.get() == menu_des5 and comboxlist3.get() == select2[menu_des5][1]:
            boxlist("department_update")
        elif comboxlist2.get() == menu_des5 and comboxlist3.get() == select2[menu_des5][2]:
            boxlist("department_enable_disable")
        elif comboxlist2.get() == menu_des5 and comboxlist3.get() == select2[menu_des5][3]:
            boxlist("department_delete")
        else:
            print("Not Exist")

    def log_clear():
        """
        日志清理：按照时间段进行清理
        :return:
        """
        initial_time = str(e1.get())
        last_time = str(e2.get())
        try:
            initial_time_2_date = datetime.datetime.strptime(initial_time, "%Y-%m-%d")
            last_time_2_date = datetime.datetime.strptime(last_time, "%Y-%m-%d")
            if initial_time_2_date > last_time_2_date:
                tkinter.messagebox.showwarning(title="Error", message=time_error1)
            else:
                askyesno = tkinter.messagebox.askyesno('Log Delete', 'Are you sure to delete logs?')
                if askyesno:
                    file_list = os.listdir(log_foler)
                    num = 0
                    for file_name in file_list:
                        try:
                            file_date_search = re.search(r"-(.*).log", file_name, re.M | re.I).group()
                            file_date_str = file_date_search[1:len(file_date_search)].replace(".log", "")
                            file_date = datetime.datetime.strptime(file_date_str, "%Y-%m-%d")
                            if file_date >= initial_time_2_date and file_date <= last_time_2_date:
                                file_path = log_foler + "\\" + file_name
                                os.remove(file_path)
                                num += 1
                        except:
                            print("can not delete file " + file_path)
                    tkinter.messagebox.showinfo(title="Delete Info", message="You have deleted " + str(num) + " files!")
        except Exception as e:
            tkinter.messagebox.showerror(title="Delete Error", message=str(e))

    def back_up():
        """
        日志备份：将整个日志文件夹迁移至选中路径
        :return:
        """
        try:
            folder = filedialog.askdirectory()
            while True:
                folder = folder + "/" + "logs"
                if os.path.exists(folder):
                    folder = folder + "/" + "logs"
                else:
                    break
            if folder:
                shutil.copytree(log_foler, folder)
                tkinter.messagebox.showinfo(title="Back-up Info", message="Back-up successfully! New log foler is " + str(folder) + ".")
        except Exception as e:
            tkinter.messagebox.showerror(title="Back-up Error", message=str(e))

    def open_folder():
        global log_foler
        try:
            folder = filedialog.askdirectory().replace("/","\\")
            if folder:
                tkinter.messagebox.showinfo(title="Open Info", message="Old path: " + str(log_foler) + "\n" + "New path: " + str(folder))
                log_foler = folder
        except Exception as e:
            tkinter.messagebox.showerror(title="Open Error", message=str(e))

    def export_file(type):
        listBoxContent = mylist.get(0, tkinter.END)
        if len(listBoxContent) > 0 and listBoxContent[0] != "No Data!":
            folder = filedialog.askdirectory().replace("/", "\\")
            if folder:
                nowtime = datetime.datetime.now().replace(microsecond=0)
                timeStamp = int(nowtime.timestamp())
                if type == "txt":
                    suffix = ".txt"
                elif type == "excel":
                    suffix = ".xls"
                elif type == "word":
                    suffix = ".docx"
                file_path = folder + "\\" + str(timeStamp) + suffix
                while True:
                    if os.path.exists(file_path):
                        nowtime = nowtime + datetime.timedelta(seconds=1)
                        timeStamp = int(nowtime.timestamp())
                        file_path = folder + "\\" + str(timeStamp) + suffix
                    else:
                        break
                try:
                    if type == "txt":
                        with open(file_path, 'wb') as f:
                            f.write(bytes("\n".join(listBoxContent),"utf-8"))
                    elif type == "excel":
                        xls = xlwt.Workbook()
                        sht1 = xls.add_sheet('Sheet1')
                        for i in range(len(listBoxContent)):
                            sht1.write(i,0,listBoxContent[i])
                        xls.save(file_path)
                    elif type == "word":
                        doc2 = docx.Document()  # 创建一个Document对象
                        try: #设置字体类型及大小
                            doc2.styles["Normal"].font.name = "Arial"  # 设置全局字体
                            doc2.styles["Normal"].font.size = Pt(10)
                        except:
                            pass
                        for i in range(len(listBoxContent)):
                            if "=====" in listBoxContent[i]:
                                doc2.add_paragraph('=' * 70)
                            else:
                                doc2.add_paragraph(listBoxContent[i])
                        doc2.save(file_path)  # 保存文档
                    tkinter.messagebox.showinfo(title="Download Info", message="File path: " + str(file_path))
                except Exception as e:
                    tkinter.messagebox.showerror(title="Download Error", message=str(e))
        else:
            tkinter.messagebox.showinfo(title="Download Info", message="No data to download!")

    def copy_listbox():
        listBoxContent = mylist.get(0, tkinter.END)
        if len(listBoxContent) > 0 and listBoxContent[0] != "No Data!":
            window.clipboard_clear()
            for content in listBoxContent:
                window.clipboard_append(content)
                window.clipboard_append("\n")

    def clear_listbox():
        mylist.delete(0, END)

    def create_pic():
        if not os.path.exists('window_title.png'):
            base64_data = b'AAABAAEAJDAAAAEAIACoHAAAFgAAACgAAAAkAAAAYAAAAAEAIAAAAAAAABsAABAnAAAQJwAAAAAAAAAAAAAAAAAAAAAAAAAAAAB+fn4AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB8fHwAfn5+AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19AH19fQF9fX0IfX19CH19fQh9fX0IfX19CH19fQh9fX0IfX19CH19fQh9fX0DfX19AH19fQB9fX0AfX19AH19fQN9fX0IfX19CH19fQh9fX0IfX19CH19fQh9fX0IfX19CH19fQh9fX0BfX19AH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19EH19fYl9fX29fX19vX19fb19fX29fX19vX19fb19fX29fX19vX19fb59fX2ffX19I319fQB9fX0AfX19I319fZ99fX2+fX19vX19fb19fX29fX19vX19fb19fX29fX19vX19fb19fX2JfX19EH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19SH19ffx9fX3/fX19/319ff99fX3/fX19/319ff99fX3/fX19/319ff99fX3/fX19eH19fQB9fX0AfX19eH19ff99fX3/fX19/319ff99fX3/fX19/319ff99fX3/fX19/319ff99fX38fX19SH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19J319fc19fX30fX198319ffR/f3/9gICA/4CAgP+AgID/gICA/4CAgP+AgID/gICAhYCAgACAgIAAgICAhYCAgP+AgID/gICA/4CAgP+AgID/gICA/39/f/19fX30fX198319ffR9fX3NfX19J319fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19AH19fR19fX06fHx8OYiIiEeqqqrZra2t/62trf6tra3+ra2t/q2trf6tra3/rKysh66urgCurq4ArKysh62trf+tra3+ra2t/q2trf6tra3+ra2t/6qqqtmIiIhHfHx8OX19fTp9fX0dfX19AH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19AH19fQB9fX0Ai4uLAMfHxw67u7vOu7u7/7u7u/+7u7v/u7u7/7u7u/+7u7v/urq6iL6+vgC+vr4Aurq6iLu7u/+7u7v/u7u7/7u7u/+7u7v/u7u7/7u7u87Hx8cOi4uLAH19fQB9fX0AfX19AH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAuLi4ALi4uBG6urrOurq6/7q6uv+6urr/urq6/7q6uv+6urr/ubm5iL29vQC9vb0Aubm5iLq6uv+6urr/urq6/7q6uv+6urr/urq6/7q6us64uLgRuLi4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAnp6eAJycnACcnJwAnJycAJycnACdnZ0AuLi4ALi4uBG6urrOurq6/7q6uv+6urr/urq6/7q6uv+6urr/ubm5iL29vQC9vb0Aubm5iLq6uv+6urr/urq6/7q6uv+6urr/urq6/7q6us64uLgRuLi4AJ2dnQCcnJwAnJycAJycnACdnZ0Am5ubAAAAAACcnJwAnJycAJycnACcnJwAnJycAJycnACcnJwAt7e3ALi4uBG6urrOurq6/7q6uv+6urr/urq6/7q6uv+6urr/ubm5iL29vQC9vb0Aubm5iLq6uv+6urr/urq6/7q6uv+6urr/urq6/7q6us64uLgRt7e3AJycnACcnJwAnJycAJycnACcnJwAnJycAJycnACbm5sAnJycAJycnC2cnJyOnJycjZycnCycnJwApqamAKamphCnp6fNp6en/6enp/+np6f/p6en/6enp/+np6f/p6enhqenpwCnp6cAp6enhqenp/+np6f/p6en/6enp/+np6f/p6en/6enp82mpqYQp6enAJycnACcnJwlnJyciJycnJOcnJw1nJycAJycnACcnJwAnJycGZycnMecnJz/nJyc/5ycnMScnJwVnJycAJubmxCbm5vMm5ub/5ubm/+bm5v/m5ub/5ubm/+bm5v/nJychJubmwCbm5sAnJychJubm/+bm5v/m5ub/5ubm/+bm5v/m5ub/5ubm8ybm5sQnJycAJycnBCcnJy5nJyc/5ycnP+cnJzRnJycH5ycnACcnJwAnJycRZycnPmcnJz/nJyc/5ycnPOcnJw3nJycAJycnBCcnJzMnJyc/5ycnP+cnJz/nJyc/5ycnP+cnJz/nJychJycnACcnJwAnJychJycnP+cnJz/nJyc/5ycnP+cnJz/nJyc/5ycnMycnJwQnJycAJycnDWcnJzxnJyc/5ycnP+cnJz6nJycSJycnACcnJwAnJycS5ycnPqcnJz/nJyc/5ycnPScnJw7nJycAJycnBCcnJzMnJyc/5ycnP+cnJz/nJyc/5ycnP+cnJz/nJychJycnACcnJwAnJychJycnP+cnJz/nJyc/5ycnP+cnJz/nJyc/5ycnMycnJwQnJycAJycnDucnJz0nJyc/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyc/5ycnPOcnJw7nJycAJ2dnQ+cnJzMnJyc/5ycnP+cnJz/nJyc/5ycnP+cnJz/nJychJORjgCTkY4AnJychJycnP+cnJz/nJyc/5ycnP+cnJz/nJyc/5ycnMydnZ0PnJycAJycnDucnJzznJyc/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyc/5ycnPOcnJ06eHFlAJSTkBGZmZjNmpmZ/5qZmf6amZn+mpmZ/pqZmf6amZn/mJeWiCoOAAUqDgAFmJeWiJqZmf+amZn+mpmZ/pqZmf6amZn+mpmZ/5mZmM2Uk5AReHFlAJycnTqcnJzznJyc/5ycnP+cnJz6nJycS5ycnACbm5sAnJycS5SUlPqIiIj/iIiI/5WVlfOamZg9Qi4PSUYzFqpYSDHwW003/VtNN/1bTTf9W003/VtNN/1cTTf+UkEp3UQwE7pEMBO6UkEp3VxNN/5bTTf9W003/VtNN/1bTTf9W003/VhIMfBGMxaqQi4PSZqZmD2VlZXziIiI/4iIiP+UlJT6nJycS5ubmwCMjIwAjY2NS4GBgfp8fHz/fHx8/4ODg/FlWUmNQi4P7EIuEP9BLQ7/QSwO/0EsDv9BLA7/QSwO/0EsDv9BLA7/QS0P/0IuEP9CLhD/QS0P/0EsDv9BLA7/QSwO/0EsDv9BLA7/QSwO/0EtDv9CLhD/Qi4P7GVZSYyDg4PxfHx8/3x8fP+BgYH6jY2NS4yMjAB+fn4AgICAS319ffp9fX3/fX19/3x8e/lTRC7sQS0O/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/QS0P/0EtDv9BLQ7/QS0P/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/QS0O/1NELux8fHv5fX19/319ff99fX36gICAS35+fgCOjo4Aj4+PS4KCgvp8fHz/fX19/4B/fv9URS7/QS0O/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9BLQ//Szkd/1VEK/9VRCr/Sjcb/0EtD/9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/QS0O/1RFLv+Af37/fX19/3x8fP+CgoL6j4+PS46OjgCcnJwAnJycS5WVlfqKior/i4uL/5KSkP9XRzD+QCwN/0IuD/9CLg//Qi4P/0IuD/9CLg//Qi0P/0AsDv9cTDT/n5qT/7CurP+wrqv/m5aO/1dHLv9ALA7/Qi4P/0IuD/9CLg//Qi4P/0IuD/9CLg//QCwN/1dHMP6SkpD/i4uL/4qKiv+VlZX6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyd/5mYlv9cTTf+RzQX/0g1Gf9INRn/SDUZ/0g1Gf9INRn/SDUZ/2BROv+lop3/vLy9/7u7u/+7u7v/vLy9/6Gdl/9cTDX/SDUZ/0g1Gf9INRn/SDUZ/0g1Gf9INRn/RzQX/1xNN/6ZmJb/nJyd/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyc/5ycnP6fnJf9oJyW/6Cclv+gnJb/oJyW/6Cclv+gnJb/oZ2X/7CurP+7u7v/urq6/7q6uv+6urr/urq6/7u7u/+vrar/oZyW/6Cclv+gnJb/oJyW/6Cclv+gnJb/oJyW/5+cl/2cnJz+nJyc/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyc/52dnf61tbX8vLy9/7y8vP+8vLz/vLy8/7y8vP+8vLz/vLy8/7u7u/+6urr/urq6/7q6uv+6urr/urq6/7q6uv+7u7v/u7y8/7u8vP+7vLz/u7y8/7u8vP+8vLz/vLy9/7W1tfydnZ3+nJyc/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqbm5v/m5ub/56env6zs7P8urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+9vb3/vb29/729vf+9vb3/vb29/729vf+6urr/urq6/7Ozs/yenp7+m5ub/5ubm/+cnJz6nJycS5ycnACbm5sAm5ubS5GRkfqEhIT/hYWF/5SUlP6zs7P8urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/8DAwP/d3d3/4uLi/+Hh4f/h4eH/4uLi/9nZ2f+8vLz/urq6/7Ozs/yUlJT+hYWF/4SEhP+RkZH6m5ubS5ubmwCJiYkAioqKS4CAgPp8fHz/fHx8/4ODg/6wsLD8u7u7/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/76+vv/T09P/1tbW/9XV1f/V1dX/1tbW/9DQ0P+8vLz/u7u7/7CwsPyDg4P+fHx8/3x8fP+AgID6ioqKS4mJiQB+fn4Afn5+Rn19ffl9fX3/fX19/4CAgP6tra38u7u7/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/u7u7/62trfyAgID+fX19/319ff99fX35fn5+Rn5+fgB/f38AgICAGn5+fsd9fX3/fX19/4iIiP+ysrL/u7u7/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/u7u7/7Kysv+IiIj/fX19/319ff9+fn7GgICAGn9/fwBra2sAe3t7AICAgCd+fn55g4ODiJeXl4a0tLSFurq6h7q6uoa6urqGurq6hrq6uoa6urqGurq6iK+vr9SsrKz8rKys+qysrPqsrKz6rKys+qysrPyvr6/Uurq6iLq6uoa6urqGurq6hrq6uoa6urqGurq6h7S0tIWXl5eGgoKCiH5+fnmAgIAne3t7AGxsbACGhoYAbGxsAH9/fwB7e3sAc3NzAIuLiwC0tLQAurq6ALq6ugC6uroAurq6ALq6ugC4uLgA////AJubm6qcnJz/nJyc/5ycnP+cnJz/nJyc/5ycnP+bm5uq////ALi4uAC6uroAurq6ALq6ugC6uroAurq6ALS0tACLi4sAc3NzAHt7ewB/f38AbW1tAIaGhgAAAAAAqqqqAIiIiACJiYkAlJSUAJ+fnwC7u7sAurq6ALq6ugC8vL0AiIB0ADsmBgBBLQ8AAAAAAJ6en6qdnZ7/nZ6e/52env+dnp7/nZ6e/52dnv+enp+qAAAAAEEtDwA7JgYAiIB0ALy8vQC6uroAurq6ALu7uwCfn58AlJSUAIiIiACIiIgAqqqqAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALq6ugC6uroAurq6Arq6ujm8vb14ioJ2hEQxE4VFMRSFRTIVh2ZaSNZvZlf/b2VW/m9lVv5vZVb+b2VW/m9mV/9mWkjWRTIVh0UxFIVEMROFioJ2hLy9vXi6uro6urq6Arq6ugC6uroAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALq6ugC6uroAurq6Xbq6uuu8vL3/iYF1/0IuD/9CLhD/Qi4Q/0EtD/9BLA7/QS0O/0EtDv9BLQ7/QS0O/0EsDv9BLQ//Qi4Q/0IuEP9CLg//iYF1/7y8vf+6urrrurq6Xbq6ugC6uroAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAJKSkgCcnJwAnJycALGxsQC7u7sVurq60Lq6uv+8vL3/iYF1/0ItD/9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLQ//iYF1/7y8vf+6urr/urq60Lu7uxWxsbEAnJycAJycnACSkpIAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwAmZmZAJCQkAe1tbU0urq667q6uv+8vL3/iYF1/0ItD/9CLhD/QS0P/0QxE/9FMRT/QS0P/0IuEP9CLhD/QS0P/0UxFP9EMRP/QS0P/0IuEP9CLQ//iYF1/7y8vf+6urr/urq667W1tTSQkJAHmpqaAJycnACcnJwAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwEnJycbZycnMCioqLNuLi4+bq6uv+8vL3/iYF1/0ItD/9BLQ//VUQp/6ujmP+wqZ//W0ox/0EtDv9BLQ7/W0ox/7Cpn/+ro5j/VUQp/0EtD/9CLQ//iYF1/7y8vf+6urr/uLi4+aKios2cnJzAnJycbZycnAScnJwAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwfnJyc4JycnP+hoaH/uLi4/rq6uv+8vL3/iYF1/0ItD/8/Kwz/hHdl/+zt7f/u7/D/kIZ2/z8rDf9AKw3/kIZ2/+7v8P/s7e3/g3dl/z8rDP9CLQ//iYF1/7y8vf+6urr/uLi4/qGhof+cnJz/nJyc4JycnB+cnJwAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwmnJyc55ycnP+hoaH9uLi4/rq6uv+8vL3/iYF1/0EtD/9ALA7/ZlY//87Lxv/T0Mz/b2BK/0AsDv9ALA7/b2BK/9PQzP/Oy8b/ZlY//0AsDv9BLQ//iYF1/7y8vf+6urr/uLi4/qGhof2cnJz/nJyc6Jubmyebm5sAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwZnJyc1pycnP+hoaH/uLi4/rq6uv+8vL3/jod8/0IuEP9CLhD/Qi4Q/1ZFKv9YRy3/Qy8R/0IuEP9CLhD/Qy8R/1hHLf9WRCr/Qi4Q/0IuEP9CLhD/jod8/7y8vf+6urr/uLi4/qGhof+cnJz/mpqa55KSkiaSkpIAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwBnJycSZubm5OkpKSpubm59rq6uv+7u7z/qKWg/1JBJ/9AKw3/QS0P/0AsDf9ALA3/QS0P/0EtD/9BLQ//QS0P/0AsDf9ALA3/QS0P/0ArDf9SQSf/qKWg/7u7vP+6urr/ubm59qSkpKeUlJTGh4eH535+fiV/f38AAAAAAAAAAAAAAAAAAAAAAJycnACcnJwAnJycAK+vrwC7u7srurq667q6uv+6urr/urq6/5eRiP9cTTb/TDof/0w5Hv9MOR7/TDke/0w5Hv9MOR7/TDke/0w5Hv9MOR7/TDof/1xNNf+XkYj/urq6/7q6uv+6urr/urq668HBwSd7e3tvfX196H19fSV9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAACcnJwAnJycALq6ugC6urokurq65bq6uv+6urr/urq6/7u7u/+ysa//p6Sf/6ainf+mop3/pqKd/6ainf+mop3/pqKd/6ainf+mop3/p6Sf/7Kxr/+7u7v/urq6/7q6uv+6urr/urq65cLCwiB8fHxwfX196H19fSV9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALq6ugC6uroHurq6pbq6uv+6urr/urq6/7q6uv+7u7v/u7u8/7u7vP+7u7z/u7u8/7u7vP+7u7z/u7u8/7u7vP+7u7z/u7u8/7u7u/+6urr/urq6/7q6uv+6urr/urq6pf///wN9fX1yfX196H19fSV9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALq6ugC6uroAurq6Ibq6uqG6urrfurq65bq6uuW6urrlurq65bq6uuW6urrlurq65bq6uuW6urrlurq65bq6uuW6urrlurq65bq6uuW6urrlurq65bq6ut+6urqhurq6IZCQkAB9fX10fX197H19fSZ9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALu7uwC6uroAurq6ALq6ugW6uroeurq6JLq6uiS6urokurq6JLq6uiS6urokurq6JLq6uiS6urokurq6JLq6uiS6urokurq6JLq6uiS6urokurq6JLq6uh66uroFurq6AHx8fAB9fX1UfX19qn19fRt9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6AH19fQB9fX0DfX19B319fQF9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAH19fQB9fX0AfX19AH19fQB9fX0AAAAAAAAAAADgAAAAcAAAAOAAAABwAAAA4AAAAHAAAADgAAAAcAAAAOAAAABwAAAA4AAAAHAAAADgAAAAcAAAAP4AAAfwAAAAgAAAABAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAIAAAAAQAAAA+AAAAfAAAAD4AAAB8AAAAMAAAAAwAAAAwAAAADAAAADAAAAAMAAAAMAAAAAwAAAAwAAAADAAAADAAAAAMAAAAMAAAAAwAAAAwAAAADAAAADgAAAAMAAAAPgAAAAwAAAA+AAAADAAAAD4AAAAMAAAAPwAAAAwAAAA////+DAAAAA='
            img_data = base64.b64decode(base64_data)
            # 注意：如果是"data:image/jpg:base64,"，那你保存的就要以png格式，如果是"data:image/png:base64,"那你保存的时候就以jpg格式。
            with open('window_title.png', 'wb') as f:
                f.write(img_data)

    def delete_pic():
        try:
            os.remove("window_title.png")
        except Exception as e:
            print(e)

    window = tk.Tk()
    # 解析log程序路径
    file_path = os.path.abspath(__file__)
    while True:
        try:
            file_path = file_path[0:file_path.rfind("\\")]
            if os.path.exists(file_path + "\\" + "logs"):
                log_foler = file_path + "\\" + "logs"
                break
        except:
            break
    if not log_foler:
        window.withdraw()
        tkinter.messagebox.showwarning(title="Error", message=path_error)
        delete_pic()
        window.destroy()
    else:
        #更改窗体图标
        create_pic()
        window.iconbitmap("window_title.png")
        delete_pic()
        # 窗口名
        window.title('Log Parser')
        window.geometry("200x200+-10+0")
        # 设定窗口大小（长*宽）
        width = 960
        height = 500
        screenwidth = window.winfo_screenwidth()
        screenheight = window.winfo_screenheight()
        alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
        window.geometry(alignstr)

        # title
        menubar = tk.Menu(window)
        # 创建一个File菜单项（默认不下拉，下拉内容包括New，Open，Save，Exit功能项）
        log_menu = tk.Menu(menubar, tearoff=0)

        # 将上面定义的空菜单命名为log_menu，放在菜单栏中，就是装入那个容器中
        menubar.add_cascade(label='Log', menu=log_menu)

        # 在Log中加入New、Open、Save等小菜单，即我们平时看到的下拉菜单，每一个小菜单对应命令操作.

        log_menu.add_command(label='Open', command=open_folder)
        log_menu.add_command(label='Back-up', command=back_up)
        log_menu.add_separator()  # 添加一条分隔线
        log_menu.add_command(label='Delete', command=log_clear)
        log_menu.add_separator()  # 添加一条分隔线
        #导出三级菜单
        export_menu = tk.Menu(log_menu, tearoff=0)
        log_menu.add_cascade(label='Export', menu=export_menu)
        export_menu.add_command(label='txt', command=lambda:export_file('txt'))
        export_menu.add_command(label='excel', command=lambda:export_file('excel'))
        export_menu.add_command(label='word', command=lambda:export_file('word'))
        log_menu.add_separator()  # 添加一条分隔线
        log_menu.add_command(label='Exit', command=window.quit)

        # 第7步，创建一个Edit菜单项（默认不下拉，下拉内容包括Cut，Copy，Paste功能项）
        editmenu = tk.Menu(menubar, tearoff=0)
        # 将上面定义的空菜单命名为 Edit，放在菜单栏中，就是装入那个容器中
        menubar.add_cascade(label='Edit', menu=editmenu)

        # 同样的在 Edit 中加入Copy、Clear等小命令功能单元，如果点击这些单元, 就会触发do_job的功能
        editmenu.add_command(label='Copy', command=copy_listbox)
        editmenu.add_command(label='Clear', command=clear_listbox)

        # 让主菜单显示出来
        window.config(menu=menubar)

        # Label(window, text="").grid(row=0, column=0, padx=1)
        Label(window, text=start_date_des, anchor="e", width=10).grid(row=0, column=0)
        Label(window, text=start_time_des, anchor="e", width=10).grid(row=0, column=3)
        Label(window, text=end_date_des, anchor="e", width=10).grid(row=0, column=5)
        Label(window, text=end_time_des, anchor="e", width=10).grid(row=0, column=7)
        Label(window, text=process_type, anchor="e", width=10).grid(row=1, column=3)
        Label(window, text=specific_user, anchor="e", width=10).grid(row=1, column=0)
        Label(window, text=search_type, anchor="e", width=10).grid(row=1, column=5)
        Label(window, text=chosen_type, anchor="e", width=10).grid(row=1, column=7)
        e1 = DateEntry(window,locale='en_US', date_pattern='y-mm-dd',width=12, background='CornflowerBlue',foreground='white', borderwidth=2)
        e2 = DateEntry(window,locale='en_US', date_pattern='y-mm-dd',width=12, background='CornflowerBlue',foreground='white', borderwidth=2)
        e3 = Entry(window,width=15)
        e3.insert(0,"00:00:00")
        e4 = Entry(window, width=15)
        e4.insert(0,"23:59:59")
        e5 = Entry(window, width=15)
        # 第一行第一列
        e1.grid(row=0, column=1, padx=10, pady=5)
        e2.grid(row=0, column=6, padx=10, pady=5)
        e3.grid(row=0, column=4, padx=10, pady=5)
        e4.grid(row=0, column=8, padx=10, pady=5)
        e5.grid(row=1, column=1, padx=10, pady=5)
        comvalue1 = tkinter.StringVar()
        comboxlist1 = ttk.Combobox(window, textvariable=comvalue1, width=12)
        comboxlist1.grid(row=1, column=4, padx=10, pady=5)
        comboxlist1['value'] = list(select1.keys())
        comboxlist1.configure(state="readonly")
        # comboxlist1.state('readonly')
        comvalue2 = tkinter.StringVar()
        comboxlist2 = ttk.Combobox(window, textvariable=comvalue2, width=12)
        comboxlist2.grid(row=1, column=6, padx=10, pady=5)
        comvalue3 = tkinter.StringVar()
        comboxlist3 = ttk.Combobox(window, textvariable=comvalue3, width=12)
        comboxlist3.grid(row=1, column=8, padx=10, pady=5)

        def xFunc1(even):
            comboxlist2.delete(0, END)
            value = comvalue1.get()
            if value:
                comboxlist2['value'] = select1[value]
                comboxlist2.current(0)  # 设置默认值
                if value == process_des1:
                    comboxlist3["value"] = ("",)
                    comboxlist3.current(0)
                    comboxlist3.configure(state="disabled")
                elif value == process_des2:
                    value2 = comvalue2.get()
                    comboxlist3['value'] = select2[value2]
                    comboxlist3.current(0)
                    comboxlist3.configure(state="readonly")
            else:
                comboxlist3.set("")
                comboxlist2['value'] = ""
                comboxlist3['value'] = ""

        def xFunc2(even):
            comboxlist3.delete(0, END)
            val = comvalue1.get()
            value = comvalue2.get()
            if val == process_des2 and value == select1[val][5]:
                comboxlist3["value"] = ""
                comboxlist3.set("")
            else:
                comboxlist3['value'] = select2[value]
                comboxlist3.current(0)

        comboxlist1.bind("<<ComboboxSelected>>", xFunc1)
        comboxlist2.bind("<<ComboboxSelected>>", xFunc2)
        # options_tuple1 = ("流程化模块","非流程化模块")
        # options_tuple1_1 = ("新建流程", "更改流程", "删除流程", "运行流程", "放行流程", "定时流程", "指挥流程", "清理日志")
        # options_tuple1_2 = ("用户菜单","角色菜单","组织菜单","公司菜单","部门菜单")
        # comboxlist2["values"] = ("新建流程", "更改流程", "删除流程", "运行流程", "放行流程", "定时流程", "指挥流程", "清理日志")
        # comboxlist2.current(0)  # 选择第一个
        # comboxlist2['state'] = 'readonly'
        # 点击事件
        # comboxlist.bind("<<ComboboxSelected>>", boxlist)

        # 如果表格大于组件，那么可以使用sticky选项来设置组件的位置
        # 同样你需要使用N，E，S,W以及他们的组合NE，SE，SW，NW来表示方位
        # Label(window, text="").grid(row=0, column=, padx=20)
        Button(window, text=btn1_des, width=10, command=getting_click).grid(row=0, column=9, sticky=W, padx=10, pady=5)
        Button(window, text=btn2_des, width=10, command=to_exit).grid(row=1, column=9, sticky=E, padx=10, pady=5)

        # 创建一个水平滚动条
        scrollbarl = Scrollbar(window, orient=HORIZONTAL)
        # 水平滚动条位于窗口底端，当窗口改变大小时会在X方向填满窗口
        # Label(window, text="").grid(row=3, column=0, padx=10)
        # scrollbarl.grid(row=3, column=1, sticky=S, pady=310)
        scrollbarl.place(relx=0.025, rely=0.97, relheight=0.025, relwidth=0.95)
        # 创建一个垂直滚动条
        scrollbar2 = Scrollbar(window)
        # 垂直滚动条位于窗口右端，当窗口改变大小时会在Y方向填满窗口
        # scrollbar2.grid(row=2, column=0, sticky=E)
        scrollbar2.place(relx=0.98, rely=0.1, relheight=0.9, relwidth=0.018)
        # 创建一个列表框， x方向的滚动条指 令是scrollbarl 对象的set()方法，
        # y方向的滚动条指令是scrollbar2对象的set()方法
        mylist = Listbox(window, xscrollcommand=scrollbarl.set, yscrollcommand=scrollbar2.set, height=10, width=100, bd=2, foreground="Navy", selectbackground="WhiteSmoke", selectforeground="Red")# 10/11/2021
        # 获取合同.docx中的数据
        # text_data = text_contract()
        # data_list = text_data.split('\n')
        # for data in data_list:
        #     mylist.insert(END, data)
        # 列表框位于窗口左端，当窗口改变大小时会在X与Y方向填满窗口
        mylist.place(relx=0.5, rely=0.55, relheight=0.75, relwidth=0.92, anchor=CENTER)
        #mylist.grid(row=2, column=1, rowspan=1, columnspan=100)
        # 移动水平滚动条时,改变列表框的x方向可见范围
        scrollbarl.config(command=mylist.xview)
        # 移动垂直滚动条时，改变列表框的y方向可见范围
        scrollbar2.config(command=mylist.yview)
        # 主窗口循环
        window.mainloop()

if __name__ == '__main__':
    pass
    log_window()
